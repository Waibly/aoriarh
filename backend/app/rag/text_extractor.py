import io

import docx
import pymupdf4llm


class TextExtractor:
    """Extracts text from PDF, DOCX, and TXT files and returns Markdown."""

    def extract(self, file_bytes: bytes, file_format: str) -> str:
        extractors = {
            "pdf": self._extract_pdf,
            "docx": self._extract_docx,
            "txt": self._extract_txt,
        }
        extractor = extractors.get(file_format)
        if not extractor:
            raise ValueError(f"Format non supporté : {file_format}")
        return extractor(file_bytes)

    # Above this threshold (in bytes), use fast text extraction instead of layout analysis
    _LARGE_PDF_THRESHOLD = 5 * 1024 * 1024  # 5 MB

    def _extract_pdf(self, file_bytes: bytes) -> str:
        import pymupdf

        doc = pymupdf.open(stream=file_bytes, filetype="pdf")

        if len(file_bytes) > self._LARGE_PDF_THRESHOLD:
            # Fast extraction for large PDFs — skip expensive layout analysis
            pages = [page.get_text("text") for page in doc]
            doc.close()
            return "\n\n".join(pages)

        md_text = pymupdf4llm.to_markdown(doc)
        doc.close()
        return md_text

    def _extract_docx(self, file_bytes: bytes) -> str:
        doc = docx.Document(io.BytesIO(file_bytes))
        parts: list[str] = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1]

            if tag == "p":
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break
                if para is None:
                    continue

                text = para.text.strip()
                if not text:
                    continue

                style = (para.style.name or "").lower()
                if "heading 1" in style:
                    parts.append(f"# {text}")
                elif "heading 2" in style:
                    parts.append(f"## {text}")
                elif "heading 3" in style:
                    parts.append(f"### {text}")
                elif "heading" in style:
                    parts.append(f"#### {text}")
                elif "list" in style:
                    parts.append(f"- {text}")
                else:
                    parts.append(text)

            elif tag == "tbl":
                for table in doc.tables:
                    if table._element is element:
                        parts.append(self._table_to_markdown(table))
                        break

        return "\n\n".join(parts)

    def _table_to_markdown(self, table: docx.table.Table) -> str:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        lines: list[str] = []
        # Header row
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
        # Data rows
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")

        return "\n".join(lines)

    def _extract_txt(self, file_bytes: bytes) -> str:
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                return file_bytes.decode(encoding)
            except (UnicodeDecodeError, ValueError):
                continue
        return file_bytes.decode("utf-8", errors="replace")
